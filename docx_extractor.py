"""
docx_extractor.py
-----------------
Extracts text and layout metadata from a .docx file.

Outputs the SAME contract as extractor.extract_pdf() so the rest of the
pipeline (segmenter, entity_extractor, skill_enricher, assembler) works
identically for both PDF and DOCX inputs.

Output:
{
    "raw_text": str,
    "lines": [
        {
            "text": str,
            "font_size": float,
            "is_bold": bool,
            "page": int        # always 0 (DOCX has no page concept)
        }
    ],
    "pages": int               # always 1 for DOCX
}

How DOCX metadata maps to line metadata
----------------------------------------
- font_size : paragraph run font size, or heading level proxy if no explicit size
- is_bold   : any run in the paragraph is bold, OR paragraph style is a Heading
- Headings  : Word heading styles (Heading 1/2/3) are treated as bold + large font
              so the segmenter's existing scoring logic picks them up correctly
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
import re


# Proxy font sizes for heading levels when no explicit size is set.
# These are chosen to sit above the typical body text median (~11pt)
# so the segmenter's "relatively larger" signal fires correctly.
_HEADING_SIZE = {
    1: 18.0,
    2: 14.0,
    3: 12.0,
    4: 11.5,
    5: 11.0,
    6: 11.0,
}
_BODY_SIZE_DEFAULT = 11.0


def _paragraph_font_size(para) -> float:
    """Return the effective font size for a paragraph in points."""
    # Check runs for explicit size
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt

    # Check paragraph style for size
    style = para.style
    if style and style.font and style.font.size:
        return style.font.size.pt

    # Heading level proxy
    style_name = (style.name or "").lower() if style else ""
    for level in range(1, 7):
        if f"heading {level}" in style_name:
            return _HEADING_SIZE[level]

    return _BODY_SIZE_DEFAULT


def _paragraph_is_bold(para) -> bool:
    """Return True if the paragraph is bold (any run bold, or heading style)."""
    # Heading styles are always treated as bold
    style_name = (para.style.name or "").lower() if para.style else ""
    if "heading" in style_name:
        return True

    # Check runs
    for run in para.runs:
        if run.bold:
            return True

    # Check paragraph-level bold
    if para.runs and all(r.bold is None for r in para.runs):
        # Inherited — check style
        style = para.style
        while style:
            if style.font and style.font.bold:
                return True
            style = style.base_style if hasattr(style, "base_style") else None

    return False


def _clean_text(text: str) -> str:
    """Clean up common DOCX text artifacts."""
    text = text.replace("\u00a0", " ")   # non-breaking space
    text = text.replace("\u2019", "'")   # right single quote
    text = text.replace("\u2018", "'")   # left single quote
    text = text.replace("\u201c", '"')   # left double quote
    text = text.replace("\u201d", '"')   # right double quote
    text = text.replace("\u2013", "–")   # en dash
    text = text.replace("\u2014", "—")   # em dash
    text = text.replace("\u2022", "•")   # bullet
    text = text.replace("\uf0b7", "•")   # wingdings bullet
    text = re.sub(r"\s+", " ", text)
    return text.strip()


# Label prefixes commonly found in template-based resumes
_LABEL_PREFIX = re.compile(
    r"^(name|email|phone|mobile|contact|address|linkedin|github|location|"
    r"visa status|current location|skype)\s*:\s*",
    re.IGNORECASE
)


def _strip_label_prefix(text: str) -> str:
    """Strip label prefixes like 'Name: ', 'Email: ' from contact lines."""
    return _LABEL_PREFIX.sub("", text).strip()


def extract_docx(path: str) -> dict:
    """
    Extract text and layout metadata from a .docx file.

    Parameters
    ----------
    path : str
        Absolute or relative path to the .docx file.

    Returns
    -------
    dict with keys: raw_text, lines, pages, has_tables
        has_tables: bool — True if the document uses tables for layout
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    doc = Document(str(path))
    lines = []

    for para in doc.paragraphs:
        text = _clean_text(para.text)
        if not text:
            continue

        # Strip label prefixes from contact-style lines
        text = _strip_label_prefix(text)
        if not text:
            continue

        font_size = _paragraph_font_size(para)
        is_bold   = _paragraph_is_bold(para)

        lines.append({
            "text":      text,
            "font_size": round(font_size, 2),
            "is_bold":   is_bold,
            "page":      0,
        })

    # Extract table content — skills tables (Category | Items) are common
    # Each row becomes a "Category: item1, item2" line so the skills parser handles it
    has_tables = len(doc.tables) > 0
    for table in doc.tables:
        # Detect if this looks like a two-column skills table
        is_skills_table = (
            len(table.columns) == 2
            and len(table.rows) >= 2
        )
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if not cells:
                continue

            if is_skills_table and len(cells) == 2:
                # Format as "Category: items" so skills parser picks it up
                text = f"{cells[0]}: {cells[1]}"
            else:
                text = "  |  ".join(cells)

            if text:
                lines.append({
                    "text":      text,
                    "font_size": _BODY_SIZE_DEFAULT,
                    "is_bold":   False,
                    "page":      0,
                })

    # Apply same split-header merge as PDF extractor
    # DOCX rarely needs this but some exported templates have letter-spaced headings
    from extractor import _merge_split_headers
    lines = _merge_split_headers(lines)

    return {
        "raw_text":  "\n".join(l["text"] for l in lines),
        "lines":     lines,
        "pages":     1,
        "has_tables": has_tables,
    }


# ── quick test ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python docx_extractor.py <file.docx>")
        sys.exit(1)

    result = extract_docx(sys.argv[1])
    print(f"Lines     : {len(result['lines'])}")
    print(f"Raw chars : {len(result['raw_text'])}")
    print()
    print(f"{'FONT':>6}  {'BOLD':>4}  TEXT")
    print("-" * 70)
    for line in result["lines"]:
        bold_flag = "✓" if line["is_bold"] else ""
        print(f"{line['font_size']:>6.1f}  {bold_flag:>4}  {line['text'][:80]}")
